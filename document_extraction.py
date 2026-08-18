import io
import json
import re

import anthropic
import pdfplumber
from docx import Document

MAX_DOCUMENT_CHARS = 20000
DOB_PATTERN = re.compile(r"^\d{4}-\d{2}-\d{2}$")


class ExtractionError(Exception):
    pass


def extract_text(filename, file_bytes):
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext == "pdf":
        return _extract_pdf_text(file_bytes)
    if ext == "docx":
        return _extract_docx_text(file_bytes)
    raise ExtractionError("Only PDF and Word (.docx) files are supported.")


def _extract_pdf_text(file_bytes):
    parts = []
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                if page_text:
                    parts.append(page_text)
    except Exception as exc:
        raise ExtractionError("Could not read that PDF file. It may be corrupted or password-protected.") from exc
    text = "\n".join(parts).strip()
    if not text:
        raise ExtractionError("Could not find any readable text in that PDF (it may be a scanned image).")
    return text


def _extract_docx_text(file_bytes):
    try:
        document = Document(io.BytesIO(file_bytes))
    except Exception as exc:
        raise ExtractionError("Could not read that Word file. It may be corrupted or in an unsupported format.") from exc
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    text = "\n".join(parts).strip()
    if not text:
        raise ExtractionError("Could not find any readable text in that document.")
    return text


def _catalog_block(options, key):
    items = options.get(key) or []
    lines = [f'- {item["id"]}: {item["label"]}' for item in items if item.get("id") and item.get("label")]
    return "\n".join(lines) if lines else "(none available)"


def _build_prompt(document_text, options, doc_type):
    doc_label = "Initial Assessment" if doc_type == "initial_assessment" else "Reassessment"
    catalog_prompt = (
        "AVAILABLE MALADAPTIVE BEHAVIORS (use only these ids):\n"
        f"{_catalog_block(options, 'maladaptive_behaviors')}\n\n"
        "AVAILABLE REPLACEMENT PROGRAMS / SKILLS (use only these ids):\n"
        f"{_catalog_block(options, 'replacement_programs')}\n\n"
        "AVAILABLE INTERVENTION STRATEGIES (use only these ids):\n"
        f"{_catalog_block(options, 'intervention_strategies')}"
    )
    truncated = document_text[:MAX_DOCUMENT_CHARS]
    return f"""You are extracting structured client information from a clinical ABA {doc_label} report to help pre-fill a client profile form. Read the document text below and return ONLY a single JSON object (no prose, no markdown fences) with this exact shape:

{{
  "name": "client's full name, or empty string if not found",
  "dob": "date of birth in YYYY-MM-DD format if determinable, else empty string",
  "age": "age as a plain string (e.g. \\"7\\") if stated or derivable, else empty string",
  "guardian_name": "primary caregiver/guardian/parent full name, or empty string",
  "bcba_name": "the supervising or authoring BCBA's full name, or empty string",
  "maladaptive_behaviors": [{{"id": "<id from the allowed list below>", "topography": "brief plain-language description of how this behavior looks for this specific client, drawn from the document"}}],
  "replacement_programs": ["<id from the allowed list below>"],
  "intervention_strategies": ["<id from the allowed list below>"]
}}

Only use ids from the allowed lists below, never invent new ids, and only include an item if the document text actually supports it. If nothing in a category matches, return an empty list for it.

{catalog_prompt}

DOCUMENT TEXT:
{truncated}"""


def extract_client_info(api_key, document_text, options, doc_type):
    if not api_key:
        raise ExtractionError("AI document extraction is not configured yet.")

    prompt = _build_prompt(document_text, options, doc_type)

    client = anthropic.Anthropic(api_key=api_key)
    try:
        response = client.messages.create(
            model="claude-sonnet-5",
            max_tokens=2000,
            messages=[{"role": "user", "content": prompt}],
        )
    except anthropic.APIError as exc:
        raise ExtractionError(f"The AI extraction request failed: {exc}") from exc

    raw = "".join(block.text for block in response.content if getattr(block, "type", None) == "text").strip()
    raw = re.sub(r"^```(?:json)?\s*|\s*```$", "", raw).strip()

    try:
        parsed = json.loads(raw)
    except json.JSONDecodeError as exc:
        raise ExtractionError("The AI response could not be understood. Please try again.") from exc

    valid_behavior_ids = {item["id"] for item in (options.get("maladaptive_behaviors") or [])}
    valid_program_ids = {item["id"] for item in (options.get("replacement_programs") or [])}
    valid_intervention_ids = {item["id"] for item in (options.get("intervention_strategies") or [])}

    behaviors, topographies = [], {}
    for entry in parsed.get("maladaptive_behaviors") or []:
        if not isinstance(entry, dict):
            continue
        bid = entry.get("id")
        if bid in valid_behavior_ids:
            behaviors.append(bid)
            topo = (entry.get("topography") or "").strip()
            if topo:
                topographies[bid] = topo

    programs = [pid for pid in (parsed.get("replacement_programs") or []) if pid in valid_program_ids]
    interventions = [iid for iid in (parsed.get("intervention_strategies") or []) if iid in valid_intervention_ids]

    dob = (parsed.get("dob") or "").strip()
    if not DOB_PATTERN.match(dob):
        dob = ""

    return {
        "name": (parsed.get("name") or "").strip(),
        "dob": dob,
        "age": str(parsed.get("age") or "").strip(),
        "guardian_name": (parsed.get("guardian_name") or "").strip(),
        "bcba_name": (parsed.get("bcba_name") or "").strip(),
        "maladaptive_behaviors": behaviors,
        "behavior_topographies": topographies,
        "replacement_programs": programs,
        "intervention_strategies": interventions,
    }
